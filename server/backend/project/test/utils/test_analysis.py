import orjson as json
from project.analysis.utils import format_exchange, truncate, extract_text_from_file, chunk_exchange, chunk_file
from docx import Document
from odf.opendocument import OpenDocumentText
from odf.text import P
import io
import pytest

## run in docker container with command: pytest project/test/utils/test_analysis.py -v -s
def test_format_exchange():
    exchange = {
        "messages": [
            {"role": "user", "content": "hello"},
            {"role": "assistant", "content": "world"}
        ]
    }
    exchange_str = json.dumps(exchange)
    formatted = format_exchange(exchange_str)
    assert formatted == "[Subject]: hello\n[Psychologist]: world\n"

def test_truncate():
    message = "The black fox jumped in front of the white sheeps."
    # Test case 1: string shorter than limit
    truncated = truncate(message, 100)
    assert message == truncated

    # Test case 2: truncate at word boundary
    truncated = truncate(message, 23)
    assert truncated == "The black fox jumped in"

    # Test case 3: truncate mid-word (should remove partial word)
    message2 = "The quick brown"
    truncated = truncate(message2, 6)  # Would cut "quick" at "qu"
    assert truncated == "The"  # Should remove partial word

def test_extract_text_from_file():
    # Test .txt file
    txt_content = b"This is a test text file.\nIt has multiple lines.\nTesting 123."
    txt_filename = "test.txt"
    txt_result = extract_text_from_file(txt_content, txt_filename)
    assert txt_result == "This is a test text file.\nIt has multiple lines.\nTesting 123."

    # Test .docx file
    doc = Document()
    doc.add_paragraph("This is a test document.")
    doc.add_paragraph("It has multiple paragraphs.")
    doc.add_paragraph("Testing 123.")
    docx_bio = io.BytesIO()
    doc.save(docx_bio)
    docx_content = docx_bio.getvalue()
    docx_filename = "test.docx"
    docx_result = extract_text_from_file(docx_content, docx_filename)
    assert docx_result == "This is a test document.\n\nIt has multiple paragraphs.\n\nTesting 123."

    # Test .odt file
    odt = OpenDocumentText()
    p1 = P(text="This is a test ODT file.")
    p2 = P(text="It has multiple paragraphs.")
    p3 = P(text="Testing 123.")
    odt.text.addElement(p1)
    odt.text.addElement(p2)
    odt.text.addElement(p3)
    odt_bio = io.BytesIO()
    odt.save(odt_bio)
    odt_content = odt_bio.getvalue()
    odt_filename = "test.odt"
    odt_result = extract_text_from_file(odt_content, odt_filename)
    assert odt_result == "This is a test ODT file.\n\nIt has multiple paragraphs.\n\nTesting 123."

    # Test unsupported file type
    with pytest.raises(ValueError) as exc_info:
        extract_text_from_file(b"test content", "test.pdf")
    assert "Unsupported file type" in str(exc_info.value)

def test_chunk_exchange():
    exchanges = [
        {
            "messages": [
                # 50 characters for content, 6 for role prefix
                {"role": "user", "content": "It's pretty crazy how annoying writing tests are."}, 
                # 345 characters for content, 11 for role prefix
                {"role": "assistant", "content": "Isn't that right! Test's are super annoying. For this test, I'm going to type a lot of characters. This next section represents a ton of characters. In fact I'm just going to keep typing about characters. Crazy how many characters I can fit into one line, right? Watch how many characters can fit in this bad boy. I wonder if this is enough yet?"},
                # 77 characters for content, 6 for role prefix
                {"role": "user", "content": "Wow! That was 345 characters! I wonder how many more random words can be fit?"},
                {"role": "assistant", "content": "Isn't that right! Test's are super annoying. For this test, I'm going to type a lot of characters. This next section represents a ton of characters. In fact I'm just going to keep typing about characters. Crazy how many characters I can fit into one line, right? Watch how many characters can fit in this bad boy. I wonder if this is enough yet?"},
                # 66 characters for content, 6 for role prefix
                {"role": "user", "content": "You just said the same thing again! That's another 345 characters!"},
                {"role": "assistant", "content": "Isn't that right! Test's are super annoying. For this test, I'm going to type a lot of characters. This next section represents a ton of characters. In fact I'm just going to keep typing about characters. Crazy how many characters I can fit into one line, right? Watch how many characters can fit in this bad boy. I wonder if this is enough yet?"},
            ]
        },
        {
            "messages": [
                # 68 characters for content, 6 for role prefix
                {"role": "user", "content": "This is a new conversation. Insane how the AI just keeps going, huh?"},   
                {"role": "assistant", "content": "Isn't that right! Test's are super annoying. For this test, I'm going to type a lot of characters. This next section represents a ton of characters. In fact I'm just going to keep typing about characters. Crazy how many characters I can fit into one line, right? Watch how many characters can fit in this bad boy. I wonder if this is enough yet?"},
                # 83 characters for content, 6 for role prefix
                {"role": "user", "content": "In fact, I swear the test writer is just copy pasting the same thing over and over!"},
                {"role": "assistant", "content": "Isn't that right! Test's are super annoying. For this test, I'm going to type a lot of characters. This next section represents a ton of characters. In fact I'm just going to keep typing about characters. Crazy how many characters I can fit into one line, right? Watch how many characters can fit in this bad boy. I wonder if this is enough yet?"},
                # 403 characters for content, 6 for role prefix
                {"role": "user", "content": "I can't help it if the tester is lazy. Unfortunately, a long message is also needed here. I'm just going to keep typing and typing and typing until I can match the AI. This is how humanity is going to be in the future, just typing out endless walls of useless text that's indistinguishable from AI! In fact, I can put anything I want in here because it'll never be read anyway! That's how crazy this is!"}
            ]
        },
        {
            "messages": [
                # 44 characters for content, 6 for role prefix
                {"role": "user", "content": "To top it off, I'm going to just keep going!"},   
                # 78 characters for content, 6 for role prefix
                {"role": "user", "content": "And I'm going to split this into separate messages! And I'll just keep typing!"}, 
                # 19 characters for content, 6 for role prefix
                {"role": "user", "content": "No one can stop me!"},   
                {"role": "assistant", "content": "Isn't that right! Test's are super annoying. For this test, I'm going to type a lot of characters. This next section represents a ton of characters. In fact I'm just going to keep typing about characters. Crazy how many characters I can fit into one line, right? Watch how many characters can fit in this bad boy. I wonder if this is enough yet?"},
            ]
        },
    ]
    chunks, long = chunk_exchange(exchanges, 200)
    expected_chunks = [
        "[Subject]: It's pretty crazy how annoying writing tests are.\n[Psychologist]: Isn't that right! Test's are super annoying. For this test, I'm going to type a lot of characters. This next section represents a ton of characters. In fact I'm just going to keep typing about characters. Crazy how many characters I can fit into one line, right? Watch how many characters can fit in this bad boy. I wonder if this is enough yet?\n[Subject]: Wow! That was 345 characters! I wonder how many more random words can be fit?\n[Psychologist]: Isn't that right! Test's are super annoying. For this test, I'm going to type a lot of characters. This next section represents a ton of characters. In fact I'm just going to keep typing about characters. Crazy how many characters I can fit into one line, right? Watch how many characters can fit in this bad boy. I wonder if this is enough yet?\n",
        "[Subject]: You just said the same thing again! That's another 345 characters!\n[Psychologist]: Isn't that right! Test's are super annoying. For this test, I'm going to type a lot of characters. This next section represents a ton of characters. In fact I'm just going to keep typing about characters. Crazy how many characters I can fit into one line, right? Watch how many characters can fit in this bad boy. I wonder if this is enough yet?\n[Subject]: This is a new conversation. Insane how the AI just keeps going, huh?\n[Psychologist]: Isn't that right! Test's are super annoying. For this test, I'm going to type a lot of characters. This next section represents a ton of characters. In fact I'm just going to keep typing about characters. Crazy how many characters I can fit into one line, right? Watch how many characters can fit in this bad boy. I wonder if this is enough yet?\n",
        "[Subject]: In fact, I swear the test writer is just copy pasting the same thing over and over!\n[Psychologist]: Isn't that right! Test's are super annoying. For this test, I'm going to type a lot of characters. This next section represents a ton of characters. In fact I'm just going to keep typing about characters. Crazy how many characters I can fit into one line, right? Watch how many characters can fit in this bad boy. I wonder if this is enough yet?\n",
        "[Subject]: I can't help it if the tester is lazy. Unfortunately, a long message is also needed here. I'm just going to keep typing and typing and typing until I can match the AI. This is how humanity is going to be in the future, just typing out endless walls of useless text that's indistinguishable from AI! In fact, I can put anything I want in here because it'll never be read anyway! That's how crazy this is!\n",
        "[Subject]: To top it off, I'm going to just keep going!\n[Subject]: And I'm going to split this into separate messages! And I'll just keep typing!\n[Subject]: No one can stop me!\n[Psychologist]: Isn't that right! Test's are super annoying. For this test, I'm going to type a lot of characters. This next section represents a ton of characters. In fact I'm just going to keep typing about characters. Crazy how many characters I can fit into one line, right? Watch how many characters can fit in this bad boy. I wonder if this is enough yet?\n"
    ]

    expected_long = [False, False, False, True, False]
    assert chunks == expected_chunks
    assert long == expected_long

def test_chunk_file():
    # Test with txt file
    txt_content = b"This is paragraph one.\n\nThis is paragraph two.\n\nThis is a longer paragraph three that should definitely be included in the chunk because it has more words than the others."
    txt_filename = "test.txt"
    txt_chunks = chunk_file(txt_content, txt_filename)
    assert len(txt_chunks) > 0
    assert isinstance(txt_chunks, list)
    assert all(isinstance(chunk, str) for chunk in txt_chunks)
    assert "paragraph one" in txt_chunks[0]
    
    # Test with docx file
    doc = Document()
    doc.add_paragraph("This is the first paragraph in DOCX.")
    doc.add_paragraph("This is the second paragraph in DOCX.")
    doc.add_paragraph("This is a much longer third paragraph in DOCX that contains more words and should demonstrate how the chunking works with longer content that might need to be split into multiple chunks depending on the configuration.")
    docx_bio = io.BytesIO()
    doc.save(docx_bio)
    docx_content = docx_bio.getvalue()
    docx_filename = "test.docx"
    docx_chunks = chunk_file(docx_content, docx_filename)
    assert len(docx_chunks) > 0
    assert isinstance(docx_chunks, list)
    assert all(isinstance(chunk, str) for chunk in docx_chunks)
    assert "first paragraph" in docx_chunks[0]

    # Test with odt file
    odt = OpenDocumentText()
    p1 = P(text="First paragraph in ODT format.")
    p2 = P(text="Second paragraph in ODT format.")
    p3 = P(text="This is a much longer third paragraph in ODT format that contains more words and should demonstrate how the chunking works with longer content that might need to be split into multiple chunks.")
    odt.text.addElement(p1)
    odt.text.addElement(p2)
    odt.text.addElement(p3)
    odt_bio = io.BytesIO()
    odt.save(odt_bio)
    odt_content = odt_bio.getvalue()
    odt_filename = "test.odt"
    odt_chunks = chunk_file(odt_content, odt_filename)
    assert len(odt_chunks) > 0
    assert isinstance(odt_chunks, list)
    assert all(isinstance(chunk, str) for chunk in odt_chunks)
    assert "First paragraph" in odt_chunks[0]

    # Test with empty file
    empty_chunks = chunk_file(b"", "empty.txt")
    assert len(empty_chunks) == 0

    # Test with unsupported file type
    with pytest.raises(ValueError) as exc_info:
        chunk_file(b"test content", "test.pdf")
    assert "Unsupported file type" in str(exc_info.value)

    # Test chunk size limits
    # Create paragraphs that are too long for one chunk (750 words)
    long_sentence = "This is a very long sentence with many characters that we will use to test chunking. " * 50
    paragraphs = [
        f"UNIQUE_START_1 {long_sentence} UNIQUE_END_1",
        f"UNIQUE_START_2 {long_sentence} UNIQUE_END_2",
        f"UNIQUE_START_3 {long_sentence} UNIQUE_END_3",
        f"UNIQUE_START_4 {long_sentence} UNIQUE_END_4"
    ]
    long_text = "\n\n".join(paragraphs)
    long_content = long_text.encode('utf-8')
    
    # Debug the input
    print("\nInput text stats:")
    for i, para in enumerate(paragraphs):
        print(f"Paragraph {i+1} length: {len(para)} chars, {len(para.split())} words")
    
    long_chunks = chunk_file(long_content, "long.txt")
    
    # Debug output
    print("\nOutput chunk stats:")
    for i, chunk in enumerate(long_chunks):
        print(f"Chunk {i+1}:")
        print(f"  - Characters: {len(chunk)}")
        print(f"  - Words: {len(chunk.split())}")
        print(f"  - First 100 chars: {chunk[:100]}")
    
    # Verify chunking behavior
    assert len(long_chunks) > 1, "Text should be split into multiple chunks"
    
    # Check chunk sizes
    for chunk in long_chunks:
        word_count = len(chunk.split())
        print(f"Chunk size: {word_count} words")
        assert word_count <= 750, f"Chunk too large: {word_count} words"
    
    # Verify content preservation
    full_text = " ".join(long_chunks)
    # Check start markers
    assert "UNIQUE_START_1" in full_text, "First paragraph start marker missing"
    assert "UNIQUE_START_2" in full_text, "Second paragraph start marker missing"
    assert "UNIQUE_START_3" in full_text, "Third paragraph start marker missing"
    assert "UNIQUE_START_4" in full_text, "Fourth paragraph start marker missing"
    
    # Check end markers
    assert "UNIQUE_END_1" in full_text, "First paragraph end marker missing"
    assert "UNIQUE_END_2" in full_text, "Second paragraph end marker missing"
    assert "UNIQUE_END_3" in full_text, "Third paragraph end marker missing"
    assert "UNIQUE_END_4" in full_text, "Fourth paragraph end marker missing"
    
    # Check relative positions
    start1_pos = full_text.find("UNIQUE_START_1")
    end1_pos = full_text.find("UNIQUE_END_1")
    start2_pos = full_text.find("UNIQUE_START_2")
    assert start1_pos < end1_pos, "First paragraph markers out of order"
    assert end1_pos < start2_pos, "First and second paragraphs out of order"